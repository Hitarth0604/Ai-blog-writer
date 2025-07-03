# output_utils.py
from docx import Document

def blog_to_docx(title, meta, tags, body):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Meta Description: {meta}")
    doc.add_paragraph(f"Tags: {tags}")
    doc.add_paragraph("---")

    for line in body.split("\n"):
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.replace("**", ""), level=2)
        else:
            doc.add_paragraph(line.strip())

    return doc
